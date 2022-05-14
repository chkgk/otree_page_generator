import docx
from .token import get_token, merge_runs_with_same_style
from dataclasses import dataclass, field


class Experiment:
    def __init__(self, docx_file):
        self.pages = []
        self.player_fields = []
        self.group_fields = []
        self.docx_filename = docx_file
        self.docx_document = docx.Document(self.docx_filename)
        self.parse()

    def parse(self):
        # we need to add some check to make sure that the very first token defines a page
        current_page = None
        for paragraph in self.docx_document.paragraphs:
            paragraph_content = []
            merged_runs = merge_runs_with_same_style(paragraph.runs)
            for run in merged_runs:
                s = get_token(run)
                if s.type == 'page':
                    current_page = Page(title=s.title, form_model=s.form_model)
                    self.pages.append(current_page)

                if current_page is None:
                    continue

                if s.type == 'input' and s.input_type != 'button':
                    current_page.form_fields.append(s)

                if s.type == 'output' and s.var_for_template:
                    current_page.vars_for_template.append(s.var_for_template)

                paragraph_content.append(s)

            p_prefix = p_postfix = ''
            if any([e.type == 'text' for e in paragraph_content]):
                p_prefix = '<p>'
                p_postfix = '</p>'

            rendered_elements = [e.render() for e in paragraph_content if e.render() is not None]
            if rendered_elements:
                content = "".join(rendered_elements)
                current_page.paragraphs.append(f"{p_prefix}{content}{p_postfix}")

        for page in self.pages:
            if page.form_model == 'player':
                self.player_fields.append(page.render_model_fields())

            if page.form_model == 'group':
                self.group_fields.append(page.render_model_fields())

    @staticmethod
    def _render_class(class_name, fields):
        content = f"class {class_name}(Base{class_name}):\n"
        for field in fields:
            content += f"{field}\n"
        return content

    def render_player_class(self):
        return self._render_class('Player', self.player_fields)

    def render_group_class(self):
        return self._render_class('Group', self.group_fields)


@dataclass
class Page:
    title: str
    paragraphs: list = field(default_factory=lambda: [])
    form_fields: list = field(default_factory=lambda: [])
    vars_for_template: list = field(default_factory=lambda: [])
    form_model: str = 'player'

    def render_model_fields(self):
        content = ''
        for field in self.form_fields:
            if field.input_type != 'generic':
                if field.input_type == 'longstring':
                    field_name = 'LongString'
                else:
                    field_name = field.input_type.capitalize()
                content += f"    {field.variable_name} = models.{field_name}Field()\n"

        return content

    def render_class(self):
        content = f"""
class {self.title}(Page):"""

        if self.form_fields:
            fields_str = ''
            for field in self.form_fields:
                fields_str += f"'{field.variable_name}', "
            fields_str = fields_str[:-2]
            content += f"""
    form_model = '{self.form_model}'
    form_fields = [{fields_str}]
"""
        if self.vars_for_template:
            vars_str = ''
            for v in self.vars_for_template:
                vars_str += f"'{v}': '',  # add your custom variable value here \n"
            vars_str = vars_str[:-2]
            content += f"""
    def vars_for_template(self):
        return {{
            {vars_str}
        }}
"""
        if not self.vars_for_template and not self.form_fields:
            content += """
    pass
"""
        return content

    def render_template(self):
        content = f"""
{{{{ block title }}}}
    { self.title }
{{{{ endblock }}}}

{{{{ block content }}}}
"""
        for paragraph in self.paragraphs:
            content += f"    {paragraph}\n"

        content += "{{ endblock }}"

        return content
