import docx
from dataclasses import dataclass
import re


@dataclass
class Experiment:
    pages: list


@dataclass
class Page:
    title: str
    paragraphs: list
    form_fields: list
    vars_for_template: list

    def render_python_class(self):
        content = f"""
class {self.title}(Page):"""

        if self.form_fields:
            fields_str = ''
            for field in self.form_fields:
                fields_str += f"'{field}', "
            fields_str = fields_str[:-2]
            content += f"""
    form_model = 'player'  # check if this might need to be 'group'
    form_fields = [{fields_str}]
"""
        if self.vars_for_template:
            vars_str = ''
            for v in self.vars_for_template:
                vars_str += f"'{v}': '',  # add your custom variable value here \n"
            vars_str = vars_str[:-2]
            content += f"""
    def vars_for_template(self):
        return {{
            {vars_str}
        }}
"""
        if not self.vars_for_template and not self.form_fields:
            content += """
    pass
"""
        return content

    def render_template(self):
        content = f"""
{{{{ block title }}}}
    { self.title }
{{{{ endblock }}}}

{{{{ block content }}}}
"""
        for paragraph in self.paragraphs:
            content += f"    {paragraph}\n"

        content += "{{ endblock }}"

        return content


class Token:
    def render(self):
        return None


class NewPage(Token):
    def __init__(self, title):
        self.type = 'page'
        self.title = title


class Comment(Token):
    def __init__(self, text):
        self.type = 'comment'
        self.text = text

    def render(self):
        return f"{{# {self.text} #}}"


class Button(Token):
    def __init__(self, text):
        self.type = 'button'
        self.text = text

    def render(self):
        if self.text == 'next':
            return "{{ next_button }}"
        else:
            return f"<button type='button' class='btn btn-primary'>{self.text}</button>"


class Output(Token):
    def __init__(self, variable):
        self.type = 'output'
        self.variable_name = variable
        self.var_for_template = None

        parts = self.variable_name.split('.')
        if len(parts) == 1:
            self.var_for_template = self.variable_name

    def render(self):
        return f"{{{{ {self.variable_name} }}}}"


class Input(Token):
    def __init__(self, input_type, variable):
        self.type = 'input'
        self.input_type = input_type
        self.variable_name = variable

    def render(self):
        return f"{{{{ formfield '{self.variable_name}' }}}}"


class Text(Token):
    def __init__(self, run):
        self.type = 'text'
        self.run = run

    def render(self):
        prefix = postfix = ''
        if self.run.bold:
            prefix += '<b>'
            postfix += '</b>'
        if self.run.italic:
            prefix += '<i>'
            postfix += '</i>'
        return f"{prefix}{self.run.text}{postfix}"


def read_docx(filename):
    return docx.Document(filename)


def token_from_control_sequence(sequence):
    if not sequence:
        return None

    parts = sequence[1:-1].split(':')
    identifier, remainder = parts[0].strip(), "".join(parts[1:]).strip()

    token = None

    if identifier == 'page':
        token = NewPage(title=remainder)
    if identifier == 'comment':
        token = Comment(text=remainder)
    if identifier == 'button':
        token = Button(text=remainder)
    if identifier == 'output':
        token = Output(variable=remainder)
    if identifier in ['boolean', 'currency', 'integer', 'float', 'string', 'longstring']:
        token = Input(input_type=identifier, variable=remainder)

    return token


def has_same_style(previous_run, current_run):
    return  previous_run.bold == current_run.bold and \
            previous_run.italic == current_run.italic and \
            previous_run.font.color.rgb == current_run.font.color.rgb


def get_token(run):
    if run.text is None:
        return None

    match = re.search(r"\[[A-Za-z\d\s.,_!?:]*\]", run.text, re.MULTILINE)
    if match:
        return token_from_control_sequence(match.string)

    if match is None:
        return Text(run=run)


def main():
    doc = read_docx('Example.docx')
    exp = Experiment(pages=[])

    # we need to add some check to make sure that the very first token defines a page

    for paragraph in doc.paragraphs:
        # merge runs with same style
        if not paragraph.runs:
            continue
        # print([r.text for r in paragraph.runs])
        merged_runs = [paragraph.runs[0]]

        paragraph_content = []

        for i, run in enumerate(paragraph.runs[1:]):
            if has_same_style(paragraph.runs[i], run):
                merged_runs[-1].text += run.text
            else:
                merged_runs.append(run)
        # print([r.text for r in merged_runs])

        for run in merged_runs:
            s = get_token(run)
            if s.type == 'page':
                current_page = Page(title=s.title, paragraphs=[], form_fields=[], vars_for_template=[])
                exp.pages.append(current_page)

            if s.type == 'input':
                current_page.form_fields.append(s.variable_name) 

            if s.type == 'output' and s.var_for_template:
                current_page.vars_for_template.append(s.var_for_template)

            if s.type == 'text':
                pass

            paragraph_content.append(s)
            # print(paragraph_content)

        p_prefix = p_postfix = ''
        if any([e.type == 'text' for e in paragraph_content]):
            p_prefix = '<p>'
            p_postfix = '</p>'

        rendered_elements = [e.render() for e in paragraph_content if e.render() is not None]
        if rendered_elements:
            content = "".join(rendered_elements)
            current_page.paragraphs.append(f"{p_prefix}{content}{p_postfix}")

    for page in exp.pages:
        print("New Page:", page.title)
        print(page.render_python_class())
        print(page.render_template())


if __name__ == '__main__':
    main()
